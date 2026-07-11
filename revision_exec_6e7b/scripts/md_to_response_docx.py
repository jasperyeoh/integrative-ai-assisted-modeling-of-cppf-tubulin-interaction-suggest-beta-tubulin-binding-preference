#!/usr/bin/env python3
"""
Convert Response_to_Reviewers.md to a PLOS-style .docx, replicating the exact
formatting conventions of the original RESPONSE_TO_REVIEWERS_v2.docx:

  - Times New Roman throughout; Title / Heading 1 (## Reviewer) / Heading 2 (### Comment)
  - Reviewer quotes (> ...) as indented Normal paragraphs
  - **bold** -> bold run
  - "quoted manuscript text" -> red run (RGB C00000), matching the convention that
    quoted new/revised manuscript wording in [Change] blocks is highlighted red
  - `code` -> Courier New font run
  - *word* -> asterisks stripped, rendered plain (matches original: not italicized)
  - Markdown tables -> real Word tables
  - "---" horizontal rules are skipped (not rendered)

Usage:
    python3 md_to_response_docx.py INPUT.md OUTPUT.docx
"""
import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH

RED = RGBColor(0xC0, 0x00, 0x00)
FONT = "Times New Roman"
CODE_FONT = "Courier New"


def add_runs(paragraph, text, base_bold=False):
    """Parse inline markdown (**bold**, `code`, "quoted"->red, *stripped*) into runs."""
    # Tokenize: **bold**, `code`, "quoted text", *word*, or plain text
    pattern = re.compile(r'(\*\*.+?\*\*|`.+?`|".+?"|\*[^*]+?\*)')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos:m.start()])
            r.font.name = FONT
            r.bold = base_bold
        token = m.group(0)
        if token.startswith("**") and token.endswith("**"):
            r = paragraph.add_run(token[2:-2])
            r.font.name = FONT
            r.bold = True
        elif token.startswith("`") and token.endswith("`"):
            r = paragraph.add_run(token[1:-1])
            r.font.name = CODE_FONT
            r.bold = base_bold
        elif token.startswith('"') and token.endswith('"'):
            r = paragraph.add_run(token)
            r.font.name = FONT
            r.font.color.rgb = RED
            r.bold = base_bold
        elif token.startswith("*") and token.endswith("*") and not token.startswith("**"):
            r = paragraph.add_run(token[1:-1])
            r.font.name = FONT
            r.bold = base_bold
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        r.font.name = FONT
        r.bold = base_bold


def is_table_row(line):
    return line.strip().startswith("|") and line.strip().endswith("|")


def is_table_sep(line):
    return bool(re.fullmatch(r"\|[\s\-:|]+\|", line.strip()))


def main():
    if len(sys.argv) != 3:
        print("Usage: python3 md_to_response_docx.py INPUT.md OUTPUT.docx")
        sys.exit(1)
    src, dst = sys.argv[1], sys.argv[2]

    with open(src, encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(12)

    sec = doc.sections[0]
    sec.page_width = Emu(7772400)
    sec.page_height = Emu(10058400)
    sec.left_margin = Emu(1143000)
    sec.right_margin = Emu(1143000)
    sec.top_margin = Emu(914400)
    sec.bottom_margin = Emu(914400)

    i = 0
    n = len(lines)
    title_done = False

    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        # Title (single leading #)
        if stripped.startswith("# ") and not title_done:
            doc.add_heading(stripped[2:].strip(), level=0)
            title_done = True
            i += 1
            continue

        # Heading 1: ## Reviewer #N
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
            i += 1
            continue

        # Heading 2: ### Comment N.N
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            i += 1
            continue

        # Blockquote
        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Emu(228600)
            add_runs(p, stripped[2:].strip())
            i += 1
            continue

        # Table
        if is_table_row(stripped):
            table_lines = []
            while i < n and is_table_row(lines[i].strip()):
                table_lines.append(lines[i].strip())
                i += 1
            # Drop separator row (---|---)
            rows = [l for l in table_lines if not is_table_sep(l)]
            cells_matrix = [[c.strip() for c in r.strip("|").split("|")] for r in rows]
            n_cols = len(cells_matrix[0])
            table = doc.add_table(rows=len(cells_matrix), cols=n_cols)
            table.style = "Table Grid"
            for ri, row in enumerate(cells_matrix):
                for ci, cell_text in enumerate(row):
                    cell = table.cell(ri, ci)
                    cell.text = ""
                    p = cell.paragraphs[0]
                    # Strip markdown bold in table headers/cells (e.g. **DeltaTOTAL**)
                    clean = re.sub(r"\*\*(.+?)\*\*", r"\1", cell_text)
                    add_runs(p, clean, base_bold=(ri == 0))
            continue

        # Standalone bold label paragraph, e.g. **[Response]** or **Manuscript:** value
        # General paragraph (possibly spanning to blank line, but our .md keeps one
        # logical line per paragraph already)
        p = doc.add_paragraph()
        add_runs(p, stripped)
        i += 1

    doc.save(dst)
    print(f"Wrote {dst}")
    print(f"Paragraphs: {len(doc.paragraphs)}  Tables: {len(doc.tables)}")


if __name__ == "__main__":
    main()
